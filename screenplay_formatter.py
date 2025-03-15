import re
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, ns
import os
from enum import Enum



def read_parameters_from_txt(file_path):
    parameters = {}

    with open(file_path, "r") as file:
        for line in file:
            if ": " in line:
                key, value = line.strip().split(": ", 1)
                parameters[key] = value

    return parameters



def set_margins(doc: DocumentType, left_inch=1.5, right_inch=1, top_inch=1, bottom_inch=1):
    """Sets the margins of the document."""
    section = doc.sections[0]
    section.left_margin = Inches(left_inch)
    section.right_margin = Inches(right_inch)
    section.top_margin = Inches(top_inch)
    section.bottom_margin = Inches(bottom_inch)
    # Remove any column layout:
    section._sectPr.remove(section._sectPr.xpath('./w:cols')[0])

class ParagraphType(Enum): #inheritance
        SCENE=1
        CHARACTER=2
        PARENTHETICAL=3
        DIALOGUE=4
        ACTION=5
        EMPTY=6
        UNKNOWN=7


def is_scene_heading(text: str):
    return text.isupper() and text.startswith('OBRAZ')
  
def is_character_name(text):
    if len(text) == 0 or is_scene_heading(text):
        return False
    first_word: str = text.split()[0]
    return len(first_word) > 1 and all(c.isalnum() and c.isupper() for c in first_word)

def is_empty(text, last_type:ParagraphType):
    return len(text)==0 

def is_action(text, last_type):
    return not is_character_name(text) and not text.startswith('(')

def is_dialogue(text, last_type):
    return not is_character_name(text) and not is_parenthetical(text, last_type) and last_type in [ParagraphType.CHARACTER, ParagraphType.PARENTHETICAL, ParagraphType.DIALOGUE]

def is_parenthetical(text, last_type:ParagraphType):   
    return last_type==ParagraphType.CHARACTER and text.startswith('(')



def check_paragraph_type(text:str, last_type:ParagraphType) -> ParagraphType:
    if is_character_name(text):
        return ParagraphType.CHARACTER
    if is_empty(text, last_type):
        return ParagraphType.EMPTY
    if is_parenthetical(text, last_type):
        return ParagraphType.PARENTHETICAL
    if is_scene_heading(text):
        return ParagraphType.SCENE
    if is_dialogue(text, last_type):
        return ParagraphType.DIALOGUE
    if is_action(text, last_type):
        return ParagraphType.ACTION
    return ParagraphType.UNKNOWN


def remove_section_breaks(doc: DocumentType) -> DocumentType:
    if len(doc.sections) <= 1:
        return doc
    
    # Remove all sections except the last one.
    for section in doc.sections[:-1]:
        parent = section._sectPr.getparent()
        parent.remove(section._sectPr)
            
    return doc

def find_start_paragraph(doc,start_keyword):
    # Finds the index of the first scene heading in the document.
    for i, paragraph in enumerate(doc.paragraphs):
        if start_keyword in paragraph.text:
            return i
    return -1


def insert_paragraph_after(paragraph, text=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    # if style is not None:
    #     new_para.style = style
    return new_para


def format_paragraph(paragraph, font_name, font_size, line_spacing, params, last_paragraph_type):
    char_indent_left = float(params.get("Character Indent Left", 4.2))- 1.5 
    char_indent_right = float(params.get("Character Indent Right", 1))- 1
    action_indent_left = float(params.get("Action Indent Left", 1.5)) - 1.5 
    action_indent_right = float(params.get("Action Indent Right", 1)) - 1
    scene_indent_left = float(params.get("Scene Indent Left", 1.5)) - 1.5 
    scene_indent_right = float(params.get("Scene Indent Right", 1)) - 1
    dialogue_indent_left = float(params.get("Dialogue Indent Left", 2.9))- 1.5 
    dialogue_indent_right = float(params.get("Dialogue Indent Right", 2.3))- 1
    parenthetical_indent_left = float(params.get("Parenthetical Indent Left", 3.6))- 1.5 
    parenthetical_indent_right = float(params.get("Parenthetical Indent Right", 2.9))- 1

    cleaned_text = re.sub(r'\s+', ' ', paragraph.text.strip())
    paragraph.text = cleaned_text

    for run in paragraph.runs:
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)

    paragraph.paragraph_format.line_spacing_rule = 1
    paragraph.paragraph_format.line_spacing = Pt(line_spacing)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = None
    paragraph.alignment = 0
    paragraph.paragraph_format.right_indent = Inches(0)

    last_paragraph_type = check_paragraph_type(cleaned_text, last_paragraph_type)

    match last_paragraph_type:
        case ParagraphType.CHARACTER:
            paragraph.paragraph_format.left_indent = Inches(char_indent_left)
            paragraph.paragraph_format.right_indent = Inches(char_indent_right)

            if not paragraph.text.isupper():
                words = paragraph.text.split()
                split_index = 0
                for i, word in enumerate(words):
                    if not word.isupper() or (word.isupper() and len(word) == 1):
                        split_index = i
                        break

                if split_index > 0:
                    character_name = " ".join(words[:split_index])
                    non_character_name = " ".join(words[split_index:])
                    paragraph.text = character_name
                    new_paragraph = insert_paragraph_after(paragraph, non_character_name)
                    last_paragraph_type = format_paragraph(new_paragraph, font_name, font_size, line_spacing, params, last_paragraph_type)

        case ParagraphType.ACTION:
            paragraph.paragraph_format.left_indent = Inches(action_indent_left)
            paragraph.paragraph_format.right_indent = Inches(action_indent_right)
        case ParagraphType.SCENE:
            paragraph.paragraph_format.left_indent = Inches(scene_indent_left)
            paragraph.paragraph_format.right_indent = Inches(scene_indent_right)
            for run in paragraph.runs:
                run.text = run.text.upper()
        case ParagraphType.DIALOGUE:
            paragraph.paragraph_format.left_indent = Inches(dialogue_indent_left)
            paragraph.paragraph_format.right_indent = Inches(dialogue_indent_right)
        case ParagraphType.PARENTHETICAL:
            paragraph.paragraph_format.left_indent = Inches(parenthetical_indent_left)
            paragraph.paragraph_format.right_indent = Inches(parenthetical_indent_right)

    return last_paragraph_type

def format_text(doc: DocumentType, start_paragraph, font_name, font_size, line_spacing, params):
    last_paragraph_type = ParagraphType.UNKNOWN
    last_paragraph_empty = False
    paragraphs_to_delete = []

    for i, paragraph in enumerate(doc.paragraphs):
        if i < start_paragraph:
            continue

        last_paragraph_type = format_paragraph(paragraph, font_name, font_size, line_spacing, params, last_paragraph_type)

        if paragraph.text == "":
            if last_paragraph_empty:
                paragraphs_to_delete.append(paragraph)
            last_paragraph_empty = True
        else:
            last_paragraph_empty = False


    # Remove marked paragraphs, -docx does not support direct deletion of paragraphs. Instead,
    # we must work with the XML structure and manually remove the paragraph from the document.  
    for paragraph in paragraphs_to_delete:
        p = paragraph._element
        p.getparent().remove(p)



def add_page_numbers(doc):
    """Adds page numbers in the top right corner, skipping the first page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True  
    
    header = section.header
    paragraph = header.add_paragraph()
    
    paragraph.alignment = 2  
    paragraph.paragraph_format.space_after = Pt(0)
    
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    run.add_text(".")

    section.header_distance = Inches(0.5)

def format_word_file(input_path, output_path, param_file):
    params = read_parameters_from_txt(param_file)

    start_keyword = params.get("Start Formatting From", "OBRAZ 1")
    font_name = params.get("Font", "Courier")
    font_size = int(params.get("Font Size", 12))
    line_spacing = int(params.get("Line Spacing", 22))

    """Formats an existing Word file and saves it."""
    if not os.path.exists(input_path):
        print(f"Error: File '{input_path}' not found!")
        return  

    doc = Document(input_path)  
    start_paragraph = find_start_paragraph(doc, start_keyword)
    doc = remove_section_breaks(doc)
    set_margins(doc)  
    format_text(doc, start_paragraph, font_name, font_size, line_spacing, params)  
    add_page_numbers(doc)  
    doc.save(output_path)  
    print(f"Formatted file saved as: {output_path}")

if __name__ == "__main__":
    input_file = r"C:\Users\maria\Sciptron\test2.docx"
    param_file = r"C:\\Users\\maria\\Sciptron\\parameters.txt"
    output_file = r"C:\Users\maria\Sciptron\formatted_output.docx"
    
    format_word_file(input_file, output_file, param_file)
   


os.startfile(output_file)