import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement, ns
import os
from enum import Enum

print("python-docx is installed correctly!")

def set_margins(doc, left_inch=1.5, right_inch=1, top_inch=1, bottom_inch=1):
    """Sets the margins of the document."""
    section = doc.sections[0]
    section.left_margin = Inches(left_inch)
    section.right_margin = Inches(right_inch)
    section.top_margin = Inches(top_inch)
    section.bottom_margin = Inches(bottom_inch)

class ParagraphType(Enum): #inheritance
        SCHENE=1
        CHARACTER=2
        PARENTHETICAL=3
        DIALOGUE=4
        ACTION=5
        EMPTY=6
        UNKNOWN=7


def is_scene_heading(text):
    """Detects if a line is a scene heading based on screenplay format."""
    return bool(re.match(r'^(\w+\s\d+\s)?(INT\.|EXT\.)\s?[\w\s\-\–ÁÉÍÓÚÝČĎĚŇŘŠŤŮŽáéíóúýčďěňřšťůž]+(DAY|NIGHT|ráno|večer|LATER|CONTINUOUS)\.?$', text, re.IGNORECASE))
  
def is_character_name(text):
    """Determines if a line is a character name."""
    return text.isupper() and len(text.split()) <= 3

def is_empty(text, last_type:ParagraphType):
    return len(text)==0 

def is_action(text, last_type):
    return not is_character_name(text) and not text.startswith('(')

def is_dialogue(text, last_type):
    return not is_character_name(text) and not is_parenthetical(text, last_type) and last_type in [ParagraphType.CHARACTER, ParagraphType.PARENTHETICAL, ParagraphType.DIALOGUE]

def is_parenthetical(text, last_type:ParagraphType):   
    return last_type==ParagraphType.CHARACTER and text.startswith('(')



def check_paragraph_type(text:str, last_type:ParagraphType) -> ParagraphType:
    if is_character_name(text):
        return ParagraphType.CHARACTER
    if is_empty(text, last_type):
        return ParagraphType.EMPTY
    if is_parenthetical(text, last_type):
        return ParagraphType.PARENTHETICAL
    if is_scene_heading(text):
        return ParagraphType.SCHENE
    if is_dialogue(text, last_type):
        return ParagraphType.DIALOGUE
    if is_action(text, last_type):
        return ParagraphType.ACTION
    return ParagraphType.UNKNOWN


def format_text(doc):
    last_paragraph_type=ParagraphType.UNKNOWN
    last_paragraph_empty = False
    paragraphs_to_delete = [] 
    """Formats character names, sets font, and ensures 55 lines per page."""
    for paragraph in doc.paragraphs:
        cleaned_text = re.sub(r'\s+', ' ', paragraph.text.strip()) # unnecessary spaces will be removed, and any multiple spaces between words will be reduced
        paragraph.text = cleaned_text   # Assign the cleaned text back to the paragraph

         # Set font and size
        for run in paragraph.runs:
            font = run.font
            font.name = "Courier"
            font.size = Pt(12)  # 12pt font size

        # Ensure single spacing for 55 lines per page
        paragraph.paragraph_format.line_spacing_rule = 1
        paragraph.paragraph_format.line_spacing = Pt(22)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent=None
        paragraph.alignment = 0  
        paragraph.paragraph_format.right_indent = Inches(0)     


        last_paragraph_type = check_paragraph_type(cleaned_text, last_paragraph_type)

        match last_paragraph_type:
            case ParagraphType.CHARACTER:
                paragraph.paragraph_format.left_indent = Inches(2.2)
            case ParagraphType.ACTION:
                paragraph.paragraph_format.left_indent = Inches(0)
            case ParagraphType.SCHENE:
                format_scene_heading(paragraph)
            case ParagraphType.DIALOGUE:
                paragraph.paragraph_format.left_indent = Inches(1)
            case ParagraphType.PARENTHETICAL:
                paragraph.paragraph_format.left_indent = Inches(1.6) 

                
    
        if cleaned_text == "":
            if last_paragraph_empty:  # previous also non empty
                paragraphs_to_delete.append(paragraph) # marked for deleting
            last_paragraph_empty = True    

        else:
            last_paragraph_empty = False  # we find non-empty row    


    # Remove marked paragraphs, -docx does not support direct deletion of paragraphs. Instead,
    # we must work with the XML structure and manually remove the paragraph from the document.  
    for paragraph in paragraphs_to_delete:
        p = paragraph._element
        p.getparent().remove(p)       




def format_scene_heading(paragraph):
    """Formats scene headings in ALL CAPS, left-aligned"""

    paragraph.alignment = 0  # Left-align (no indent)
    paragraph.paragraph_format.left_indent = Inches(0)  # No indentation
    
    for run in paragraph.runs:
        run.text = run.text.upper()  # Ensure uppercase

# def format_parethetical(doc):              
#     character_name = None 
#     dialogue_pattern = re.compile(r'^(\w+)(\s*\(.*?\))$')

#     for paragraph in doc.paragraphs:
#         text = paragraph.text.strip()

#            # Detect character names 
#         if is_character_name(text):
#             character_name = text

#         elif character_name and text.startswith('(') and text.endswith(')'):
#             paragraph.paragraph_format.left_indent = Inches(1.6)

def add_page_numbers(doc):
    """Adds page numbers in the top right corner, skipping the first page."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True  
    
    header = section.header
    paragraph = header.add_paragraph()
    
    paragraph.alignment = 2  
    paragraph.paragraph_format.space_after = Pt(0)
    
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    run.add_text(".")

    section.header_distance = Inches(0.5)

def format_word_file(input_path, output_path):
    """Formats an existing Word file and saves it."""
    if not os.path.exists(input_path):
        print(f"Error: File '{input_path}' not found!")
        return  

    doc = Document(input_path)  
    set_margins(doc)  
    format_text(doc)  
    # format_scene_headings(doc)  
    # format_parethetical(doc)
    add_page_numbers(doc)  
    doc.save(output_path)  
    print(f"Formatted file saved as: {output_path}")

if __name__ == "__main__":
    input_file = r"C:\Users\maria\wordy\Test2.docx"
    output_file = r"C:\Users\maria\wordy\formatted_output.docx"
    
    format_word_file(input_file, output_file)


os.startfile(output_file)